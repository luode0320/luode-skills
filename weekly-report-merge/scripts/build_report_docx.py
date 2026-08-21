#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从 report.json 生成最终周报 docx（固定排版样式）。

用法:
    python build_report_docx.py report.json out.docx

report.json 结构（UTF-8）:
{
  "title": "周报（2026.08.17 星期一 - 2026.08.21 星期五）",
  "subtitle": "姓名 · 公司 · 部门",
  "intro": [["结论：", "..."], ["影响：", "..."]],        // 结论摘要八项
  "overview_headers": ["事项", "本周关键进展"],
  "overview_rows": [["...", "..."], ...],
  "progress_title": "当前开发待发布版本总体进度：",
  "progress_items": ["项目: 未完成事项 -> 项目完成度: xx%", ...],
  "sections": [{"title": "1. xxx", "bullets": ["...", ...]}, ...],
  "projects": [{"name": "xxx", "items": ["修复: ...", "进行中: ..."]}, ...],
  "plan_headers": ["#", "计划事项", "所属", "截止时间"],
  "plan_rows": [["1", "...", "兑换业务", "月底（08.31）"], ...],
  "execution_appendix": ["...", ...],
  "tracking_appendix": ["...", ...]
}

样式：正文微软雅黑/Calibri 10.5pt，标题三级色阶（1F4E79/2E74B5），表格 Table Grid。
"""
import json
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_CN = "微软雅黑"


def set_cn(run):
    run.font.name = "Calibri"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)


def build(data, out):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    for hname, sz, col in [
        ("Heading 1", 18, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12, "2E74B5"),
    ]:
        st = doc.styles[hname]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(col)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    def para(text="", bold=False, size=None, style=None):
        p = doc.add_paragraph(style=style)
        r = p.add_run(text)
        r.bold = bold
        if size:
            r.font.size = Pt(size)
        set_cn(r)
        return p

    def bullets(items):
        for it in items:
            p = doc.add_paragraph(style="List Bullet")
            set_cn(p.add_run(it))

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            set_cn(r)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                set_cn(cells[i].paragraphs[0].add_run(v))
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)

    title = doc.add_heading(data["title"], level=0)
    for r in title.runs:
        set_cn(r)
    if data.get("subtitle"):
        p = para(data["subtitle"], size=9)
        p.runs[0].font.color.rgb = RGBColor.from_string("808080")
    doc.add_paragraph()

    if data.get("intro"):
        doc.add_heading("结论摘要", level=2)
        for label, text in data["intro"]:
            p = doc.add_paragraph()
            r1 = p.add_run(label)
            r1.bold = True
            set_cn(r1)
            set_cn(p.add_run(text))

    if data.get("overview_rows"):
        doc.add_heading("一、本周概览", level=2)
        table(data["overview_headers"], data["overview_rows"], widths=[4.2, 12.0])

    if data.get("progress_items"):
        doc.add_heading("二、总体归纳统计", level=2)
        para(data.get("progress_title", "当前开发待发布版本总体进度："), bold=True)
        bullets(data["progress_items"])

    if data.get("sections"):
        doc.add_heading("三、本周工作展开", level=2)
        for sec in data["sections"]:
            doc.add_heading(sec["title"], level=3)
            bullets(sec["bullets"])

    if data.get("projects"):
        doc.add_heading("四、各项目明细（Git 提交维度）", level=2)
        p = para("以下内容来自本周 Git 提交与工作区未提交改动，作者精确匹配本人，低价值提交已过滤。", size=9)
        p.runs[0].font.color.rgb = RGBColor.from_string("808080")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            r = p.add_run(proj["name"] + ":")
            r.bold = True
            set_cn(r)
            p.paragraph_format.space_before = Pt(6)
            para("报告内容点:")
            bullets(proj["items"])

    if data.get("plan_rows"):
        doc.add_heading("五、下一步计划", level=2)
        table(data["plan_headers"], data["plan_rows"], widths=[1.0, 7.5, 2.6, 4.1])

    if data.get("execution_appendix"):
        doc.add_heading("执行附录", level=2)
        bullets(data["execution_appendix"])
    if data.get("tracking_appendix"):
        doc.add_heading("追踪附录", level=2)
        bullets(data["tracking_appendix"])

    doc.save(out)
    print("SAVED:", out)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit("用法: python build_report_docx.py report.json out.docx")
    with open(sys.argv[1], encoding="utf-8") as f:
        data = json.load(f)
    build(data, sys.argv[2])
